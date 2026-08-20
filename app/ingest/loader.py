from __future__ import annotations

import csv
import hashlib
import io
from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from pathlib import Path

# Declared accepted formats. Anything outside this list is rejected at intake with
# a named reason rather than silently producing an empty document.
ACCEPTED = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".eml": "message/rfc822",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class UnsupportedFormat(ValueError):
    """The extension is outside the declared set."""


class UnreadableDocument(ValueError):
    """Right extension, unusable content: corrupt, encrypted, empty, image-only.

    Separate from UnsupportedFormat on purpose. "We do not accept .pptx" and
    "this .pdf is damaged" are different facts, and telling someone which one
    they are looking at is the difference between a fixable problem and a
    mystery.
    """


@dataclass
class LoadedDocument:
    filename: str
    mime: str
    text: str
    sha256: str


def sha256_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()


def load_bytes(filename: str, blob: bytes) -> LoadedDocument:
    ext = Path(filename).suffix.lower()
    if ext not in ACCEPTED:
        raise UnsupportedFormat(
            f"{filename}: '{ext}' is outside the declared format set "
            f"({', '.join(sorted(ACCEPTED))}). Rejected at intake."
        )
    if not blob:
        raise UnreadableDocument(f"{filename}: the file is empty (0 bytes).")
    mime = ACCEPTED[ext]
    if ext in (".txt", ".md"):
        text = blob.decode("utf-8", errors="replace")
    elif ext == ".csv":
        text = _csv_to_text(blob)
    elif ext == ".eml":
        text = _eml_to_text(blob)
    elif ext == ".pdf":
        text = _guard(filename, "PDF", _pdf_to_text, blob)
    elif ext == ".docx":
        text = _guard(filename, "Word document", _docx_to_text, blob)
    else:  # pragma: no cover - guarded by ACCEPTED
        raise UnsupportedFormat(filename)

    if not text.strip():
        raise UnreadableDocument(
            f"{filename}: parsed successfully but contains no readable text."
        )
    return LoadedDocument(
        filename=filename, mime=mime, text=_normalise(text), sha256=sha256_bytes(blob)
    )


def load_path(p: Path) -> LoadedDocument:
    return load_bytes(p.name, p.read_bytes())


# NUL and the other C0 controls are stripped because Postgres text columns reject
# them outright. A single 0x00 anywhere in a decoded file used to surface as a 500
# from the database layer, several steps away from the file that caused it.


# Postgres text columns reject NUL outright, and the other C0 controls are noise
# inside a document. Stripped here rather than discovered at the database, where a
# single 0x00 used to surface as a 500 several layers from the file that caused it.
#
# Built from code points rather than a regex character class on purpose: the
# escape sequence for NUL is exactly the kind of thing that survives one
# round-trip through a tool and not the next.
_CONTROL_CODES = frozenset(c for c in range(32) if c not in (9, 10)) | frozenset({127})


def _strip_controls(text: str) -> str:
    return "".join(ch for ch in text if ord(ch) not in _CONTROL_CODES)

def _guard(filename: str, kind: str, fn, blob: bytes) -> str:
    """Any parser failure becomes a named rejection.

    Closes the class rather than the case: pdfminer raises PDFSyntaxError,
    python-docx raises BadZipFile, and the next library will raise something else
    again. None of them should reach a user as "Internal Server Error".
    """
    try:
        return fn(blob)
    except UnreadableDocument:
        raise
    except Exception as exc:  # noqa: BLE001
        raise UnreadableDocument(
            f"{filename}: not a readable {kind} ({type(exc).__name__}: {exc}). "
            f"If it opens in a normal reader it may be encrypted, or damaged in transit."
        ) from exc


def _normalise(text: str) -> str:
    # One newline convention everywhere: character offsets are the provenance
    # substrate, and a CRLF file would make every span platform-dependent.
    text = text.replace(chr(13) + chr(10), chr(10)).replace(chr(13), chr(10))
    text = _strip_controls(text)
    return text.strip() + chr(10)


def _csv_to_text(blob: bytes) -> str:
    rows = list(csv.reader(io.StringIO(blob.decode("utf-8", errors="replace"))))
    if not rows:
        return ""
    header, *body = rows
    out = [" | ".join(header)]
    for r in body:
        out.append(" | ".join(r))
    return "\n".join(out)


def _eml_to_text(blob: bytes) -> str:
    msg = BytesParser(policy=policy.default).parsebytes(blob)
    parts = [
        f"From: {msg.get('From', '')}",
        f"To: {msg.get('To', '')}",
        f"Date: {msg.get('Date', '')}",
        f"Subject: {msg.get('Subject', '')}",
        "",
    ]
    body = msg.get_body(preferencelist=("plain",))
    if body is not None:
        parts.append(body.get_content())
    return "\n".join(parts)


def _pdf_to_text(blob: bytes) -> str:
    import pdfplumber  # imported lazily: most corpora never touch it

    out = []
    with pdfplumber.open(io.BytesIO(blob)) as pdf:
        for page in pdf.pages:
            out.append(page.extract_text() or "")
    text = "\n\n".join(out).strip()
    if not text:
        # An image-only PDF. Honest refusal beats an empty document that looks
        # like a document with nothing in it.
        raise UnsupportedFormat(
            "PDF has no extractable text layer (scanned?). OCR is out of the "
            "declared scope -- see README 'Declared scope'."
        )
    return text


def _docx_to_text(blob: bytes) -> str:
    import docx  # python-docx, lazily

    d = docx.Document(io.BytesIO(blob))
    out = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            out.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(out)
